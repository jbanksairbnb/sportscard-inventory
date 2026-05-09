"""Sports Collective — one-page pitch (brand-styled).

Run: python3 docs/build_pitch.py
Outputs: docs/Sports-Collective-Pitch.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

PLUM = RGBColor(0x3D, 0x1F, 0x4A)
ORANGE = RGBColor(0xE2, 0x5A, 0x1C)
INK = RGBColor(0x2C, 0x1B, 0x33)
INK_SOFT = RGBColor(0x55, 0x42, 0x60)
TEAL = RGBColor(0x2D, 0x7A, 0x6E)
RUST = RGBColor(0xC5, 0x4A, 0x2C)

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(1.4); s.bottom_margin = Cm(1.4)
    s.left_margin = Cm(1.6); s.right_margin = Cm(1.6)

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10.5)
doc.styles['Normal'].font.color.rgb = INK


def shade(cell, hex_):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_)
    tcPr.append(shd)


def borders(cell, color='3D1F4A', size='8'):
    tcPr = cell._tc.get_or_add_tcPr()
    tb = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), size); b.set(qn('w:color'), color)
        tb.append(b)
    tcPr.append(tb)


# ── Header / banner ──
banner = doc.add_table(rows=1, cols=1)
cell = banner.rows[0].cells[0]
shade(cell, '3D1F4A')
borders(cell, '3D1F4A', '4')
cell.paragraphs[0].text = ''
p = cell.paragraphs[0]
r = p.add_run('SPORTS COLLECTIVE')
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(0xE2, 0x5A, 0x1C)
p2 = cell.add_paragraph()
r = p2.add_run('Built by a collector, for the collection.')
r.italic = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0xF8, 0xEC, 0xD0)
doc.add_paragraph()


def section(title, color=ORANGE, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title.upper())
    r.bold = True; r.font.size = Pt(size); r.font.color.rgb = color
    r.font.name = 'Calibri'


def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.font.size = Pt(10.5); r.font.color.rgb = INK
    return p


def bullets(items):
    for it in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        if isinstance(it, tuple):
            r = p.add_run(it[0]); r.bold = True; r.font.size = Pt(10.5)
            p.add_run(it[1]).font.size = Pt(10.5)
        else:
            p.add_run(it).font.size = Pt(10.5)


# ── PROBLEM / SOLUTION (two-column) ──
two = doc.add_table(rows=1, cols=2)
two.autofit = True
left, right = two.rows[0].cells
borders(left, 'C8B8D0', '4'); borders(right, 'C8B8D0', '4')
shade(left, 'FBF3DD'); shade(right, 'FBF3DD')

def fillcol(cell, title, paragraphs):
    cell.paragraphs[0].text = ''
    p = cell.paragraphs[0]
    r = p.add_run(title.upper())
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = ORANGE
    for line in paragraphs:
        para = cell.add_paragraph()
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run(line)
        run.font.size = Pt(10); run.font.color.rgb = INK

fillcol(left, 'The Problem',
        ['Vintage card collecting still runs on spreadsheets, screenshots, and Facebook DMs.',
         'Sellers chase payments and tag prior bidders by hand.',
         'Buyers wade through algorithmic feeds hoping to spot a card on their want list.',
         'Existing tools (Cardbase, CollX, Card Ladder) skew modern and ignore the FB-claim-sale workflow that drives 1950s–80s vintage trade.'])
fillcol(right, 'Our Solution',
        ['One workspace for the vintage collector and small-volume FB seller.',
         'Set tracking, bulk import, scan tools, want-list matching, eBay scanning.',
         'FB-auction and claim-sale runner with bidder history — the only feature no comp owns.',
         'Marketplace surfaces listings to the exact members who need them — algorithmically aligned with their want list.'])
doc.add_paragraph()


# ── MARKET ──
section('Market opportunity')
bullets([
    ('Wedge: ', 'Card-aware Facebook auction / claim-sale tooling — no serious dedicated SaaS exists today (AuctionAnything is generic, $70+/mo).'),
    ('Adjacent comps: ', 'Cardbase ($13/mo), CollX Pro ($10/mo), Card Ladder Pro ($20/mo), MyCardPost ($9/mo flat). MySlabs charges 1–3% sale fees with no sub.'),
    ('Why we win the vintage seller: ', 'eBay 13.25% + COMC ~15% + Whatnot ~11% all leave ≥10pts of margin on the table for someone with a sub-only / thin-fee model.'),
    ('Tailwind: ', 'PSA Set Registry conditioned vintage collectors to think in checklists. We ship the same mental model as a real workspace.'),
])


# ── BUSINESS MODEL ──
section('Business model')
tbl = doc.add_table(rows=4, cols=4)
tbl.autofit = True
hdr = tbl.rows[0].cells
for i, h in enumerate(['Tier', 'Price', 'For', 'Includes']):
    hdr[i].paragraphs[0].text = ''
    p = hdr[i].paragraphs[0]
    r = p.add_run(h); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(0xF8, 0xEC, 0xD0)
    shade(hdr[i], '3D1F4A'); borders(hdr[i], '3D1F4A', '6')

rows = [
    ('Free', '$0', 'Browser / new collector', '250-card cap, manual entry, basic price lookup, set view, marketplace browsing, 1 active listing'),
    ('Collector', '$9 / mo  ($90 / yr)', 'Active set-builder', 'Unlimited inventory, bulk CSV, scan tools, set-builder, eBay want-list scan, Want List PDF, historical import'),
    ('Seller', '$24 / mo  ($240 / yr)', 'FB-claim-sale runner', 'Everything above + FB auction/claim-sale runner, bidder history, sales metrics, unlimited marketplace listings'),
]
for i, (a, b, c, d) in enumerate(rows, 1):
    cells = tbl.rows[i].cells
    cells[0].paragraphs[0].text = ''
    r = cells[0].paragraphs[0].add_run(a); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = PLUM
    cells[1].paragraphs[0].text = ''
    r = cells[1].paragraphs[0].add_run(b); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = ORANGE
    cells[2].paragraphs[0].text = c; cells[2].paragraphs[0].runs[0].font.size = Pt(10)
    cells[3].paragraphs[0].text = d; cells[3].paragraphs[0].runs[0].font.size = Pt(9.5)
    for cc in cells:
        shade(cc, 'FBF3DD'); borders(cc, 'C8B8D0', '4')

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
r = p.add_run('+ 3% seller-side marketplace fee (no buyer fee). Beats eBay 13.25%, COMC ~15%, Whatnot ~11%.')
r.italic = True; r.font.size = Pt(10); r.font.color.rgb = INK_SOFT


# ── TRACTION ──
section('Where we are today')
bullets([
    'Pilot launching this week with a hand-picked group of vintage collectors.',
    'Live: set tracking, bulk CSV / PSA-Registry import, scan tools, eBay want-list scanner, marketplace, FB auction + claim-sale tools, sales metrics, historical-transaction importer, members directory.',
    'Brand identity, member onboarding, and admin tooling shipped.',
])


# ── ASK / NEXT QUARTER ──
section('Next quarter — what unlocks the model')
bullets([
    ('Stripe Connect + 3% marketplace fee plumbing — ', 'unlocks Seller-tier monetization and the wedge claim.'),
    ('FB claim-sale runner v1 with bidder export — ', 'gate Seller-only; the only feature no comp owns.'),
    ('CDN + image-resize pipeline — ', 'pre-empts the Supabase storage-egress cost cliff before scale.'),
    ('Free-tier soft cap (250 cards) + 14-day Seller trial — ', 'calibrates conversion without burning vintage goodwill.'),
    ('Vintage set-registry parity (1950s–80s baseball) — ', 'keeps Free credible against PSA Set Registry.'),
])

# Footer
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
r = p.add_run('sports-collective.com  ·  pilot release v1.0')
r.italic = True; r.font.size = Pt(9); r.font.color.rgb = INK_SOFT

out = Path(__file__).parent / 'Sports-Collective-Pitch.docx'
doc.save(out)
print(f'Wrote {out}')
