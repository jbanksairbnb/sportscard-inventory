"""Draft Sports Collective's eBay Marketplace Insights API access application.

Run: python3 docs/build_ebay_api_application.py
Outputs: docs/Ebay-Marketplace-Insights-API-Application.docx

The .docx is structured around the questions eBay typically asks on their
Buy/Marketplace API access form. Edit values inline (company contact,
volume estimates, etc.) before submitting through the eBay developer portal.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


PLUM     = RGBColor(0x3D, 0x1F, 0x4A)
ORANGE   = RGBColor(0xE2, 0x5A, 0x1C)
INK      = RGBColor(0x2C, 0x1B, 0x33)
INK_SOFT = RGBColor(0x55, 0x42, 0x60)


def shade(cell, hex_):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_)
    tcPr.append(shd)


def borders(cell, color='3D1F4A', size='8'):
    tcPr = cell._tc.get_or_add_tcPr()
    tb = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), size); b.set(qn('w:color'), color)
        tb.append(b)
    tcPr.append(tb)


doc = Document()
for s in doc.sections:
    s.top_margin = Cm(1.8); s.bottom_margin = Cm(1.8)
    s.left_margin = Cm(2.0); s.right_margin = Cm(2.0)

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)
doc.styles['Normal'].font.color.rgb = INK


def title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = PLUM


def subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.italic = True; r.font.size = Pt(11); r.font.color.rgb = INK_SOFT
    p.paragraph_format.space_after = Pt(14)


def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = ORANGE


def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = PLUM


def body(text, italic=False, soft=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.italic = italic
    r.font.size = Pt(11)
    r.font.color.rgb = INK_SOFT if soft else INK
    return p


def bullets(items):
    for it in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        if isinstance(it, tuple):
            r = p.add_run(it[0]); r.bold = True; r.font.size = Pt(11)
            p.add_run(' ' + it[1]).font.size = Pt(11)
        else:
            p.add_run(it).font.size = Pt(11)


def kv_table(rows, fill_left='F2E4D0'):
    t = doc.add_table(rows=len(rows), cols=2)
    t.autofit = True
    for i, (k, v) in enumerate(rows):
        c0 = t.rows[i].cells[0]; c1 = t.rows[i].cells[1]
        shade(c0, fill_left); borders(c0, 'C8B8D0', '4'); borders(c1, 'C8B8D0', '4')
        c0.paragraphs[0].text = ''
        rk = c0.paragraphs[0].add_run(k); rk.bold = True; rk.font.size = Pt(10.5); rk.font.color.rgb = PLUM
        c1.paragraphs[0].text = v
        c1.paragraphs[0].runs[0].font.size = Pt(10.5)


# ─────────────────────────── COVER ───────────────────────────

title('Sports Collective')
subtitle('eBay Marketplace Insights API — Application for Production Access')

h1('Section 1 — Applicant overview')
kv_table([
    ('Application / product name', 'Sports Collective'),
    ('Production URL',              'https://sports-collective.com'),
    ('Business model',              'Subscription-based collector workspace + thin-fee marketplace (in development). Currently free during pilot.'),
    ('eBay Developer ID / App ID',  '[insert from your eBay Developer Account]'),
    ('Primary contact',             '[Your name]'),
    ('Primary contact email',       '[Your email]'),
    ('Company / sole proprietorship name',
                                    '[Legal entity name, or "Sole proprietor – [Your name]"]'),
    ('Country of operation',        'United States'),
    ('Vertical',                    'Sportscards / collectibles (vintage 1950s–1990s focus)'),
])

# ─────────────────────────── PRODUCT DESCRIPTION ───────────────────────────

h1('Section 2 — Product description')

body(
    'Sports Collective is a dedicated workspace for vintage sportscard collectors. '
    'Members manage their card inventory at the set / checklist level, track which '
    'cards they own and are still chasing, capture cost and condition, attach scans, '
    'and run their selling activity (Facebook auctions, Facebook claim sales, and an '
    'on-platform marketplace) end to end inside one tool.'
)
body(
    'The product was built to replace the spreadsheets, screenshots, and Facebook '
    'Messenger threads that vintage collectors and small-volume sellers use today. '
    'It is not a reseller or arbitrage tool — every primary feature is grounded in '
    'a real collector workflow (set tracking, want-list management, run-a-claim-sale, '
    'settle-with-buyer).'
)

h2('Where eBay data fits in')
body(
    'Two existing surfaces in the product would benefit from access to eBay '
    'Marketplace Insights data:'
)
bullets([
    ('Set-level eBay Hits scan: ', 'today this scans active eBay listings against a member’s want list using the public Browse API. Adding sold-comp data would let us show the recent realized prices for the same card on the same hit, so members can value the listing before they click through.'),
    ('Per-card Research Prices modal: ', 'collectors enter recent comparable sales, weight them, and arrive at a market value for a single card. We currently surface no pre-populated comps — the user types each one. Marketplace Insights would let us pre-fill recent eBay sold prices for the exact card (year + brand + card # + condition), with the seller able to edit / re-weight / discard each one.'),
])

# ─────────────────────────── USE CASE ───────────────────────────

h1('Section 3 — Specific use case for Marketplace Insights data')

body(
    'We are requesting access to the Marketplace Insights API specifically because '
    'the use case requires recent sold (completed) listings, not just active inventory. '
    'Card prices in this market change weekly; we cannot provide a defensible '
    'market-value estimate from active listings alone (active prices are anchoring, '
    'not realized).'
)

h2('How the data will be displayed')
bullets([
    'Inside the Research Prices modal, sold comps are shown as one row per result, listing: sale date, sale price, condition / grade, source (clearly labeled "eBay sold"), and a deep link to the original eBay item page.',
    'Each comp is editable by the user — they can drop one, change its weight, or override its condition before computing their market value.',
    'No raw API responses are exposed in the UI. We never present sold data outside the context of the user’s own card-research workflow.',
    'eBay attribution ("Source: eBay Sold Listings") is shown next to each row, and the link back to the eBay item ID is preserved.',
])

h2('What the data is not used for')
bullets([
    'We do not run an automated repricing engine, an arbitrage feed, or any kind of public pricing index.',
    'We do not redistribute eBay data to third parties.',
    'We do not display eBay sold data on any non-authenticated pages or to non-members.',
    'We do not store sold-comp data beyond the cache window required to make the user experience responsive (see Section 5).',
])

# ─────────────────────────── TECHNICAL ───────────────────────────

h1('Section 4 — Technical implementation')
kv_table([
    ('Hosting',       'Vercel (Next.js 16, edge + Node functions). Backend on Supabase (Postgres, Storage).'),
    ('Auth model',    'Each end user authenticates via Supabase Auth. No Marketplace Insights data is exposed to unauthenticated visitors.'),
    ('API access',    'OAuth 2.0 client-credentials flow against eBay’s production endpoints; tokens stored server-side, never shipped to the browser.'),
    ('Endpoints used (intended)',
                      '/buy/marketplace_insights/v1_beta/item_sales/search (primary). Browse API endpoints continue to be used for active listings.'),
    ('Search shape', 'q = "{year} {brand} #{card_number} {player}"; filter on category_ids + condition + sold-within-N-days; aspect filters on Year + Set when available.'),
    ('Rate-limit posture',
                      'All calls are server-side and aggressively cached (see Section 5). We will respect Application-level and User-level limits and back off on 429 responses with exponential retry.'),
    ('Error handling',
                      'Failed lookups degrade silently to "no comps found" in the UI. We do not retry into rate-limit storms.'),
])

# ─────────────────────────── DATA HANDLING ───────────────────────────

h1('Section 5 — Data handling, retention, and privacy')

bullets([
    ('Caching: ', 'sold-comp results for a given (card, condition) tuple are cached for 24 hours to avoid hitting eBay on every modal open. Cached rows include only the public fields needed to render the row (item id, sold price, sold date, condition, image thumbnail URL, deep link).'),
    ('Retention: ', 'cached rows are evicted after 30 days. We do not maintain a permanent local mirror of eBay sold data.'),
    ('User data: ', 'Sports Collective member data (cost, profit, notes) is stored in Supabase under owner-only RLS policies. eBay sold-comp data sits alongside it but is not user-private — it is publicly available via eBay.'),
    ('No redistribution: ', 'eBay data is never exposed to anyone other than the authenticated end user who triggered the lookup.'),
    ('Compliance: ', 'we will publish a Privacy Policy and Terms of Use before opening the marketplace publicly. The eBay API Terms of Use will be linked from the developer-facing acknowledgment in our product (UI text near the comp source label).'),
])

# ─────────────────────────── VOLUME ───────────────────────────

h1('Section 6 — Estimated request volume')

body(
    'Pilot is currently in soft-launch with a hand-picked group of vintage '
    'collectors. Best-effort projections at three growth stages:'
)

t = doc.add_table(rows=4, cols=4)
hdr = t.rows[0].cells
for i, h in enumerate(['Stage', 'Active members / mo', 'Marketplace Insights calls / day (avg)',
                       'Peak / day']):
    hdr[i].paragraphs[0].text = ''
    p = hdr[i].paragraphs[0]
    r = p.add_run(h); r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = RGBColor(0xF8, 0xEC, 0xD0)
    shade(hdr[i], '3D1F4A'); borders(hdr[i], '3D1F4A', '6')

vols = [
    ('Pilot (today)',   '50–150',    '~200',     '~600'),
    ('Growth band',     '1,000',     '~2,500',   '~8,000'),
    ('Mature',          '10,000',    '~20,000',  '~60,000'),
]
for i, row in enumerate(vols, 1):
    cells = t.rows[i].cells
    for j, v in enumerate(row):
        cells[j].paragraphs[0].text = ''
        r = cells[j].paragraphs[0].add_run(v); r.font.size = Pt(10.5); r.font.color.rgb = INK
        if j == 0: r.bold = True; r.font.color.rgb = PLUM
        shade(cells[j], 'FBF3DD'); borders(cells[j], 'C8B8D0', '4')

body(
    'Volumes are demand-driven (a member opens the Research Prices modal, or '
    'requests a set-level eBay Hits scan). The 24-hour cache materially flattens '
    'the curve — repeat opens of the same card by the same or different members '
    'do not generate new calls.',
    soft=True,
)

# ─────────────────────────── COMPLIANCE ───────────────────────────

h1('Section 7 — Compliance acknowledgment')

bullets([
    'We agree to the eBay API License Agreement and Developer Program Policies as of the date of this application.',
    'We will display the eBay attribution required for sold-comp data and preserve outbound links to the original eBay item.',
    'We will not use Marketplace Insights data to power a public pricing index, an automated repricer, or any feature that competes with eBay’s own pricing surfaces.',
    'We will rate-limit ourselves below documented quotas, cache aggressively, and degrade silently on rate-limit responses.',
    'We will respond to any policy changes or compliance reviews from eBay within 5 business days.',
])

# ─────────────────────────── CONTACT ───────────────────────────

h1('Section 8 — Contact for follow-up')

kv_table([
    ('Primary contact',       '[Your name]'),
    ('Email',                 '[Your email]'),
    ('Phone (optional)',      '[Phone]'),
    ('eBay seller account (optional, if linked)', '[Your eBay handle, if applicable]'),
    ('Best response window',  'Mon–Fri, 9am–6pm ET'),
])

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Thank you for the consideration. We are happy to provide a sandbox demo or a screen recording of the Research Prices modal on request.')
r.italic = True; r.font.size = Pt(10.5); r.font.color.rgb = INK_SOFT


# Save
out = Path(__file__).parent / 'Ebay-Marketplace-Insights-API-Application.docx'
doc.save(out)
print(f'Wrote {out}')
