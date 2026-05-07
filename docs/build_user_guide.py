"""Build the Sports Collective user guide as a .docx file.

Run: python3 docs/build_user_guide.py
Outputs: docs/Sports-Collective-User-Guide.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

# Brand colors (from site CSS variables)
PLUM = RGBColor(0x3D, 0x1F, 0x4A)
ORANGE = RGBColor(0xE2, 0x5A, 0x1C)
INK = RGBColor(0x2C, 0x1B, 0x33)
INK_SOFT = RGBColor(0x55, 0x42, 0x60)
RULE = RGBColor(0xC8, 0xB8, 0xD0)
PAPER = "FBF3DD"
CREAM = "F8ECD0"

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# Base style: body font
styles = doc.styles
body = styles['Normal']
body.font.name = 'Calibri'
body.font.size = Pt(11)
body.font.color.rgb = INK

def shade_cell(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_border(cell, color="3D1F4A", size="6"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = PLUM
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = ORANGE
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12.5)
    r.font.color.rgb = PLUM
    return p

def para(text, bold=False, italic=False, size=None, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size: r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return p

def bullets(items, style='List Bullet'):
    for item in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(2)
        if isinstance(item, str):
            p.add_run(item)
        else:
            # tuple: (lead_bold, rest)
            r = p.add_run(item[0])
            r.bold = True
            p.add_run(item[1])

def numbered(items):
    for item in items:
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(2)
        if isinstance(item, str):
            p.add_run(item)
        else:
            r = p.add_run(item[0]); r.bold = True
            p.add_run(item[1])

def callout(title, body_text, fill=CREAM):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    shade_cell(cell, fill)
    set_cell_border(cell, color="3D1F4A", size="8")
    cell.paragraphs[0].text = ''
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = PLUM
    r.font.size = Pt(11)
    p2 = cell.add_paragraph()
    r2 = p2.add_run(body_text)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = INK_SOFT
    doc.add_paragraph()

def image_placeholder(label, capture_hint, route=None):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    shade_cell(cell, "F2E4D0")
    set_cell_border(cell, color="E25A1C", size="8")
    cell.paragraphs[0].text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"[ Screenshot — {label} ]")
    r.bold = True
    r.font.color.rgb = ORANGE
    r.font.size = Pt(11)
    if route:
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(f"Capture from: {route}")
        r2.italic = True
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = INK_SOFT
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(capture_hint)
    r3.font.size = Pt(9.5)
    r3.font.color.rgb = INK_SOFT
    doc.add_paragraph()

def page_break():
    doc.add_page_break()

# ───────────────────────── COVER ─────────────────────────

cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(120)
r = cover.add_run("SPORTS COLLECTIVE")
r.bold = True
r.font.size = Pt(40)
r.font.color.rgb = ORANGE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("USER  GUIDE")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = PLUM

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
r = p.add_run("Built by a collector, for the collection.")
r.italic = True
r.font.size = Pt(13)
r.font.color.rgb = INK_SOFT

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(220)
r = p.add_run("Pilot release · v1.0")
r.font.size = Pt(10)
r.font.color.rgb = INK_SOFT

page_break()

# ───────────────────────── INTRO ─────────────────────────

h1("Welcome to Sports Collective")

para(
    "Most platforms are built for sellers. This one is built for you. We've eliminated the "
    "friction of vintage collecting — from automated set tracking that prevents duplicate "
    "purchases to streamlined tools that bypass the noise of social media. We provide the "
    "digital infrastructure your physical collection deserves."
)

h3("What you can do here")
bullets([
    ("Inventory intelligence — ", "manage your sets with tools designed for the vintage workflow. Stop buying duplicates."),
    ("Curation over algorithms — ", "we connect you directly to the inventory you actually need, instead of hoping Facebook surfaces it."),
    ("Faster Facebook posts — ", "auction and claim-sale tools that turn 30 minutes of copy-paste into a couple of clicks."),
    ("Community first — ", "a dedicated space where the love of the hobby outranks the hustle of the flip."),
])

h3("How this guide is organized")
para(
    "Sections 1–4 cover the basics every member needs: getting in, setting up your profile, "
    "and building your first set. Sections 5–8 cover selling tools — listings, the marketplace, "
    "and the Facebook auction/claim-sale workflow. Sections 9–11 cover the data, community, "
    "and admin features. Appendices at the end have keyboard shortcuts and the pilot launch "
    "checklist."
)

callout(
    "💡 Tip — Read this first",
    "If you only have ten minutes, read Sections 1, 2, and the first half of Section 3. "
    "That's enough to apply, sign in, set your profile, and start your first set. "
    "You can come back for the selling tools later."
)

page_break()

# ───────────────────────── TABLE OF CONTENTS ─────────────────────────

h1("Contents")
toc = [
    "1. Getting started",
    "2. Your profile",
    "3. Building your shelf — Sets",
    "4. Adding card images",
    "5. The want list & Default Target",
    "6. Creating listings",
    "7. The marketplace & purchases",
    "8. Facebook sales tools — auctions & claim sales",
    "9. Sales metrics & historical transactions",
    "10. Community",
    "11. Admin",
    "12. Tips & troubleshooting",
    "Appendix A — Keyboard shortcuts",
    "Appendix B — Pilot launch checklist",
]
for line in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(line)
    r.font.size = Pt(11.5)
    r.font.color.rgb = INK

page_break()

# ───────────────────────── 1. GETTING STARTED ─────────────────────────

h1("1. Getting started")

h2("Apply for an account")
numbered([
    "From the public landing page, click Sign up / Request access.",
    "Create your login credentials (email + password).",
    "Fill out the Request Access form. We use this to confirm you're a collector and not a bot — give a real answer.",
    "Submit. You'll see a Pending screen and receive a confirmation email.",
])
image_placeholder(
    "Apply / Request Access screen",
    "Open /apply in an incognito window so you're signed out, then full-page screenshot.",
    route="/apply",
)

h2("What happens after you apply")
para(
    "An admin reviews each application by hand. You'll get a second email when your account "
    "is approved — usually within 24 hours during the pilot. While you wait, your account "
    "exists but you'll be bounced to the Pending page if you try to sign in."
)

h2("Sign in")
numbered([
    "Go to the login page.",
    "Enter the email + password you used to apply.",
    "After approval, login lands you on the Home feed.",
])

h2("Forgot your password")
para(
    "Click Forgot password? on the login screen. We'll email a reset link. The link opens "
    "the Reset Password page where you can set a new one. Links expire after one hour."
)

callout(
    "Note",
    "If your reset email doesn't arrive within a couple of minutes, check spam, then ask "
    "the admin to resend. We currently send from a transactional address that some Gmail "
    "filters route aggressively."
)

page_break()

# ───────────────────────── 2. PROFILE ─────────────────────────

h1("2. Your profile")

para(
    "Your profile is how other members find you and see what you collect. The first time "
    "you sign in, your home page is intentionally sparse — fill it in so the rest of the "
    "site has something to work with."
)

h2("Photos")
bullets([
    ("Cover photo — ", "the wide banner across the top of your home and public profile. Use something landscape-oriented (1600×500 or so works well)."),
    ("Avatar (the round one) — ", "shows up next to your name everywhere on the site and on every listing you post."),
])
image_placeholder(
    "Home page with empty profile",
    "Sign in as a brand-new test user, screenshot /home before any photos or sets are added.",
    route="/home",
)

h2("Identity")
bullets([
    ("Display name — ", "what shows up on listings and posts. Use whatever you'd want a buyer to see."),
    ("Handle — ", "your @-style short name. Members can search for you by handle in the directory."),
    ("Bio — ", "a sentence or two about what you collect. Keep it short."),
    ("City — ", "optional, helps regional collectors find each other."),
])

h2("Collecting interests")
bullets([
    ("Favorite team / favorite players / what you're chasing — ", "free-text fields shown on your profile. These also feed future matchmaking — when another member lists a card on your chase list, you'll see it."),
    ("Favorite Cards (the showcase) — ", "pin a handful of crown-jewel cards from your collection. These appear at the top of your public profile."),
])
image_placeholder(
    "Filled-out profile",
    "Capture a fully-populated /home (cover, avatar, bio, favorite cards). Use your own account.",
    route="/home",
)

h2("Privacy — what others can see")
para(
    "Your profile has a Profile shared toggle. When it's on, other approved members can "
    "view your profile via the Members directory. When it's off, you only appear to admins. "
    "Either way, your set data is never made public unless you explicitly create a share link "
    "(Section 3)."
)

callout(
    "🔒 What we never share",
    "Your email address, your cost basis on cards, and any data on the Sales Metrics page "
    "are private to you. Other members only see what shows on your public profile and "
    "any sets you've explicitly shared."
)

page_break()

# ───────────────────────── 3. SETS ─────────────────────────

h1("3. Building your shelf — Sets")

para(
    "My Shelf is where every set you're tracking lives. A set is a list of cards (the "
    "checklist) plus your status on each one — owned, condition, cost, value, notes. "
    "It's the foundation everything else on the site is built on: the want list, eBay "
    "hits, Facebook sales tools, and the marketplace all read from your sets."
)
image_placeholder(
    "My Shelf overview",
    "Capture / (root) once you have 3–6 sets in progress. Show the grid of set cards with progress bars.",
    route="/",
)

h2("Create a new set")

h3("From a pre-loaded template (recommended)")
numbered([
    "Click New upload on My Shelf.",
    "Click Select a set — a dropdown of pre-loaded checklists appears.",
    "Pick the set you're building. The title pre-fills; rename it if you want (you can name it anything — many of us run multiple copies of the same set, e.g. \"1971 Topps — graded\", \"1971 Topps — high-grade build\", \"1971 Topps — for sale\").",
    "Click Save and edit set. The full checklist drops in pre-populated.",
])
image_placeholder(
    "New set picker",
    "Click New upload on / and capture the modal with the set dropdown open.",
    route="/set/new",
)

h3("From a CSV (your own checklist or a PSA Set Registry export)")
numbered([
    "On My Shelf click New upload → Upload CSV.",
    "Click Choose file and pick your CSV. PSA users: open your PSA collection and click Download CSV (right side of the page above the search box). The file is named set (##).csv — that's the one.",
    "Name the set and click Save and edit set.",
    "PSA imports populate Owned, Grade, and any values you stored on PSA. Custom CSVs follow the standard template — download the template from the upload modal if you're starting from scratch.",
])

h2("The set table")
para(
    "After saving, you land on the set's edit page. Each card is one row with columns for "
    "Owned, Grade, Cost, Value, Notes, and image. You can edit a single row inline, or "
    "select multiple rows and bulk-apply a value."
)

h3("Bulk edit")
numbered([
    "Click the checkbox at the start of any row, or use the top-row checkbox to select all.",
    "Pick the field you want to update from the bulk-edit bar (e.g. Condition).",
    "Type the value once.",
    "Click Apply to N rows. Every selected row updates in one save.",
])
para(
    "I use this constantly. Common pattern: select all, set condition to the average grade for "
    "the set, then go back through and override individual rows that are different. Same flow "
    "for Owned status — set everything to Yes, then flip the few you don't have to No."
)

h3("Default Target")
para(
    "Click Default Target at the top of the set. A modal asks what condition you're targeting "
    "for the cards you don't yet own — graded vs raw, plus a low/high range. Saving the target "
    "activates two things:"
)
bullets([
    "Set-wide eBay Hits scan on the home feed (Section 5) — only listings inside your range come back.",
    "Future matchmaking against the marketplace and Facebook auction/claim-sale tools — when another seller lists a card you need in your range, it surfaces in your feed.",
])
image_placeholder(
    "Default Target modal",
    "Open any set, click Default Target, capture the modal.",
    route="/set/[slug]",
)

h3("Show only Needed")
para(
    "There's a toggle at the top of the set table that hides everything you've marked Owned. "
    "Useful when you're shopping — flip it on, and the table becomes your live want list for "
    "this set."
)

h3("Inventory mode (lightbox)")
para(
    "Click View on a set to open the inventory view — a clean grid of the cards with images. "
    "Click any image to open the lightbox. Use ← / → keys to flip through the set, Esc to close."
)

h2("Sharing a set")
numbered([
    "On My Shelf, find the set card and click 🔗 Share.",
    "If the set isn't shared yet, the system creates a public share token, marks the set SHARED, and copies the link to your clipboard.",
    "If it's already shared, the existing link is copied and you see a ✓ Copied confirmation.",
    "Paste the link anywhere — anyone with the link can view the set in read-only mode. They cannot see your cost or notes.",
])

callout(
    "Heads-up",
    "Share links are unguessable but unauthenticated — anyone you give the link to can open it. "
    "If you want to revoke access, edit the set and toggle sharing off; the old link stops working."
)

page_break()

# ───────────────────────── 4. IMAGES ─────────────────────────

h1("4. Adding card images")

para(
    "Three ways to get card images into the system, depending on what you're doing. "
    "All three live under the My Shelf or Listings sections."
)

h2("Single-card scans (for listings)")
para(
    "Loads images one card at a time and creates a draft listing per card. Best for "
    "individual sales where each card needs its own write-up."
)
numbered([
    "Choose how your scans are organized: Fronts then backs, Interleaved F/B/F/B, or Fronts only.",
    "Drag-and-drop the files (or click to browse). You can load multiple pairs in one go.",
    "Fill out the card info for each card and Save and go to next.",
    "When you're done, click Go to my listings — your drafts are waiting in the Drafts section, ready to edit or activate.",
])
image_placeholder(
    "Scan inbox / single-card scan flow",
    "Capture /listings/scan-inbox with two image pairs loaded.",
    route="/listings/scan-inbox",
)

h2("Set-inventory scans (for filling a set)")
para(
    "Loads images straight into a set you already own. Best for scanning a finished or "
    "near-finished set in batches. Works great with sheet-fed scanners."
)
numbered([
    "Open the set you want to add images to.",
    "Select the card numbers you're about to scan — important: keep them in the same order the cards will scan.",
    "Click Start scanning.",
    "Pick the scan layout (F/B alternating, or all fronts then all backs).",
    "Drag the images into the upload zone, or click to browse.",
    "Click Save all and choose whether to replace any existing images for those cards.",
])
callout(
    "🔧 Sweet spot",
    "Batches of 25 cards / 50 images run smoothly. Bigger batches start to get sluggish in the browser."
)

h2("Multi-card scan (6-up sheets)")
para(
    "If your scanner produces a 6-card sheet (2 rows of 3), this tool slices it into "
    "individual card images for you — no Photoshop required."
)
numbered([
    "Open one of your sets.",
    "Click into the Multi-card scan tool.",
    "Upload the front-side sheet (2 rows of 3 cards).",
    "Drag the grid handles until each card is cleanly framed.",
    "Click Split into 6 cards. The six tiles appear in a row underneath.",
    "Click Next: Backs and repeat for the back-side sheet, keeping the same alignment.",
    "If a back-image ended up paired with the wrong front, drag it onto the correct slot.",
    "Click Save → View set to confirm everything landed where you wanted it.",
])
image_placeholder(
    "Multi-card scan grid editor",
    "Open /listings/scan-multi-card with a 6-card test image, capture the grid-adjustment view.",
    route="/listings/scan-multi-card",
)

page_break()

# ───────────────────────── 5. WANT LIST ─────────────────────────

h1("5. The want list & Default Target")

para(
    "Your want list is calculated automatically. Every card across every set you've marked "
    "as not Owned counts as a want. There's no separate want list to maintain — you just "
    "keep your sets honest, and the want list stays current."
)

h2("How matching works")
bullets([
    "Each set has a Default Target (condition + price range).",
    "When eBay hits, marketplace listings, or Facebook auctions come in, we filter against the matching set's target before showing them to you.",
    "Cards with no condition listed on the seller's side still come through — we'd rather show you a possible match than silently drop it.",
])

h2("Set a Default Target")
numbered([
    "Open the set.",
    "Click Default Target.",
    "Choose Graded or Raw, then a low/high condition.",
    "Set a target price range if you want price-based filtering too.",
    "Click Save Target.",
])

h2("Want-list hits feed")
para(
    "On the home feed, the Want-list hits panel shows every recent listing — from the marketplace, "
    "from Facebook auctions/claim sales posted by other members, and from priority eBay sellers — "
    "that matches a card on your want list within your target. Click through to the listing or refresh to "
    "rescan."
)
image_placeholder(
    "Home feed with Want-list hits panel",
    "On /home, scroll to Your Feed and capture the Want-list hits + eBay hits + Auction hits stack.",
    route="/home",
)

h2("eBay hits")
para(
    "A per-set eBay scanner. Pick a set from the dropdown and click Search eBay. We pull "
    "listings for every card you don't own and filter to your target condition. The scan covers "
    "general eBay search plus a list of priority sellers we follow more aggressively."
)

page_break()

# ───────────────────────── 6. LISTINGS ─────────────────────────

h1("6. Creating listings")

para(
    "Listings are cards you've decided to sell. They live in three states — Draft, Active, "
    "Sold — and they feed every other selling tool on the site (the marketplace, the Facebook "
    "auction tool, and the claim-sale tool all pull from your active listings). Until a card "
    "is active, it can't be put into an auction or a claim sale."
)
image_placeholder(
    "My Listings — three lifecycle tabs",
    "On /listings, screenshot the page with at least one card in each of Draft, Active, Sold.",
    route="/listings",
)

para(
    "From My Listings in the top nav you have three creation paths."
)

h2("Blank listing")
numbered([
    "Click New blank listing.",
    "Fill in card description, condition, asking price, your cost (optional — used only for your private profit math), free-text description, shipping options, and shipping charges.",
    "Click Create listing. The listing lands in Drafts.",
    "Click Edit on the draft to add up to five photos: click the box with the +, choose images from your computer, click Save changes.",
    "When the listing is ready, click Activate. It's now live in the marketplace and available to drop into Facebook auctions/claim sales.",
])

h2("From inventory")
para(
    "Pull cards directly out of a set you've built — useful when you have a dedicated "
    "\"for sale\" copy of a set."
)
numbered([
    "Click New listing → From inventory.",
    "Pick the source set, click Next: Pick cards.",
    "Check the cards you want to sell, click Review.",
    "Pick the shipping option once and click Apply to all rows. Fill in any per-card overrides.",
    "Click Create drafts. You land in the Drafts section, ready to activate.",
])

h2("Bulk upload (CSV)")
para(
    "The fastest way to load a big batch — multiple years, sets, or boxes at once."
)
numbered([
    "Click New listing → Bulk upload.",
    "Click Download template.",
    "Fill in the rows. Not every column is required — see the example rows in the template.",
    "Delete the example rows, save as .csv (UTF-8).",
    "Click Choose CSV file and select your file. The system shows a preview table and flags any formatting problems.",
    "Click Import listings as drafts. The drafts are created in one batch — activate them when ready.",
])

callout(
    "💡 Workflow tip",
    "Many of us keep a \"For Sale\" set per year (e.g. \"1971 Topps — singles for sale\") and "
    "use the From inventory path. The CSV path is best for one-off lots or when you're moving "
    "from a spreadsheet."
)

page_break()

# ───────────────────────── 7. MARKETPLACE ─────────────────────────

h1("7. The marketplace & purchases")

h2("How the marketplace works")
para(
    "Every active listing across every member is searchable in the marketplace. Buyers "
    "filter by player, year, brand, set, condition, or price. When a card matches a card on "
    "their want list inside their Default Target, it gets pinned with a chase badge."
)
image_placeholder(
    "Marketplace search results",
    "Capture /marketplace with at least one chase-badge listing visible (use a test buyer with a target set).",
    route="/marketplace",
)

h2("Selling on the marketplace")
para(
    "There's nothing extra to do — once a listing is active, it's in the marketplace. "
    "Buyers contact you through the listing; you finalize the sale outside the platform "
    "(the pilot doesn't process payments yet — see Section 12)."
)

h2("Search")
para(
    "The /search page is a global card search across every listing and every set headline. "
    "Type a player name, year, or set; we group results by where the match came from "
    "(active listings, your sets, other members' shared sets)."
)

h2("Purchases")
para(
    "The Purchases page is the buyer-side mirror of My Listings. It tracks every card you've "
    "bought through the platform: who you bought from, paid status, shipping status, and "
    "your running total of card spend."
)
image_placeholder(
    "Purchases page",
    "Capture /purchases for a buyer with a few completed orders.",
    route="/purchases",
)

page_break()

# ───────────────────────── 8. FACEBOOK SALES ─────────────────────────

h1("8. Facebook sales tools — auctions & claim sales")

para(
    "These tools were the reason I built this site. I was tired of running auctions out of "
    "spreadsheets, scrolling through old comments to find prior bidders to tag, and chasing "
    "payments with Messenger DMs. The Facebook tools turn an auction or claim sale into a "
    "few clicks while still posting in your group on Facebook the way buyers expect."
)

callout(
    "Before you start",
    "These tools pull from your Active listings only. Build and activate the listings you "
    "want to include before you start an auction or claim sale (see Section 6)."
)

h2("Set up your templates (one-time)")
para(
    "Templates are the boilerplate that wraps every post — shipping policy, bidding rules, "
    "payment instructions. Set them once; the system fills in the per-card details for each "
    "auction."
)
numbered([
    "From Auctions, click Templates → + New.",
    "There are three template types: Single-card auction, Multi-card auction (each card in the comments), and Winning-bid message.",
    "Anything inside {} is a placeholder we fill in automatically — don't edit those.",
    "Everything else (description, shipping, bidding rules, payment instructions) is yours to customize.",
    "Click Create template. You can edit it any time.",
])
image_placeholder(
    "Auction templates list",
    "Capture /fb-auctions/templates with at least one Single-card and one Multi-card template defined.",
    route="/fb-auctions/templates",
)

h2("Run an auction")

h3("Start it")
numbered([
    "From the home top nav click FB Sales → Auctions.",
    "Click + New auction. Choose Single card or Multi card.",
    "Pick the card(s) from your active inventory.",
    "An Edit post body panel appears below the table — tweak starting bid, minimum raise, any auction-specific notes.",
    "Click Generate auction.",
])

h3("Post it")
para(
    "We generate the post text, plus a Download side-by-side image button if your listing "
    "has a front + back (creates a single image with front on the left, back on the right). "
    "Copy the text into your Facebook group, attach the image, post."
)
numbered([
    "Click Go live.",
    "Paste the URL of your live Facebook post into the URL field — buyers (and our matchmaking) will deep-link to it later.",
])

h3("Track bids")
para(
    "On the active-auction page, log every bid you receive. Two reasons this matters:"
)
bullets([
    "We build a bidder history per card, per group. Next time you run a similar lot, the system can suggest prior bidders to tag.",
    "Sales metrics (Section 9) reads from this data — losing bids drive your cost-of-acquisition curves and help you price future lots.",
])
image_placeholder(
    "Active auction page",
    "Open an in-progress auction at /fb-auctions/[id], capture with 3–4 bids logged.",
    route="/fb-auctions/[id]",
)

h3("Close it out")
numbered([
    "When the auction ends, change status to Ended (item is won, payment pending) or Sold (paid).",
    "Click Manage. The system generates the winning-bid Messenger text — billing total, shipping, payment instructions — pre-filled and ready to paste.",
    "Mark Sold once payment hits. The auction moves to the Sold section of the Auctions home.",
])

h2("Run a claim sale")
para(
    "From the FB Sales menu pick Claim sales. Two flavours:"
)
bullets([
    ("Single card — ", "one card per comment, first claim wins."),
    ("Multi-card / group lot — ", "a group of up to six cards offered together. Buyer claims the whole lot in one comment."),
])

numbered([
    "Click + New claim sale.",
    "Click + Single card or + Group lot to add a row. Set price (per card or group total).",
    "Pick the card(s) from your active listings — make sure you've added images first.",
    "Add as many rows as you need; mix singles and group lots in one sale.",
    "Click Generate and manage. The system tells you if anything is missing.",
    "On the result page, copy the post body text into your Facebook group.",
    "Click Auto-build to assemble a 6-up collage from the front/back images you've already uploaded — front sheet and back sheet, ready to post as comments.",
])
image_placeholder(
    "Claim sale builder",
    "On /fb-claim-sales/new, capture mid-build with 4–5 cards and one group lot configured.",
    route="/fb-claim-sales/new",
)

h3("Manage the sale")
numbered([
    "Back on Claim sales home, your post shows every card with its claim status.",
    "When someone claims a card in the comments, fill in their info on the row.",
    "Click Manage on the post → Invoices to generate a Messenger-ready invoice for that buyer.",
    "When payment hits, mark the row Paid; the invoice can be marked Paid separately so you can track outstanding balances.",
])

h2("Bidder data")
para(
    "Every bidder you log builds a profile under FB Sales → Bidders. For each bidder we "
    "track the cards they've bid on, won, lost, and watched ('w' tag requests). Use this to:"
)
bullets([
    "Tag prior bidders on similar lots (the right people, not 200 random tags).",
    "Spot your top buyers and what they collect.",
    "Set realistic reserves on future auctions of similar cards.",
])
image_placeholder(
    "Bidder leaderboard",
    "Capture /fb-auctions/bidders showing top bidders by total spend and watch-list size.",
    route="/fb-auctions/bidders",
)

page_break()

# ───────────────────────── 9. SALES METRICS ─────────────────────────

h1("9. Sales metrics & historical transactions")

para(
    "Sales metrics is the analytics view across every transaction you've logged — auctions "
    "won, claim sales, marketplace orders, and historical Facebook activity you've imported. "
    "It's the answer to \"how is the side hustle actually doing?\"."
)
image_placeholder(
    "Sales metrics dashboard",
    "Capture /sales-metrics with at least three months of activity. Show the totals tiles, the channel breakdown, and the top-buyer leaderboard.",
    route="/sales-metrics",
)

h2("What you'll see")
bullets([
    "Total revenue, cost basis, and profit (revenue minus your stored cost on each card).",
    "Channel mix — what came from Facebook auctions vs. claim sales vs. the marketplace.",
    "Top buyers across all your sales.",
    "Card-by-card profit on every won sale.",
])

h2("Importing historical transactions")
para(
    "If you've been running Facebook sales before joining the platform, the Historical "
    "transactions page lets you backfill that data so your metrics aren't starting from zero."
)
numbered([
    "Open Sales metrics → Historical transactions.",
    "Click ⬇ Download Excel template. The workbook has dropdowns for engagement type and channel, plus an Instructions sheet.",
    "Fill in your transactions — each row is one sale, lost bid, or tag request.",
    "Save the file as CSV (Comma delimited). The template's Instructions sheet has step-by-step save instructions for Excel and Google Sheets.",
    "Back on the Historical transactions page, click ⬆ Upload CSV. A live progress bar runs through every row.",
    "The summary shows ✓ N added and lists any rows that failed (typically missing bidder name or a bad date format).",
])
callout(
    "Why historical data matters",
    "Bidders, groups, and won/lost ratios all roll up across historical and live data. Importing "
    "a year of past activity gives the matchmaking and bidder-tag tools real signal from day one."
)

page_break()

# ───────────────────────── 10. COMMUNITY ─────────────────────────

h1("10. Community")

h2("The Members directory")
para(
    "Members is the directory of every approved collector with a public profile. Sort by "
    "cards owned, sets tracked, want-list size, or name; search by display name, handle, "
    "city, or bio."
)
image_placeholder(
    "Members directory",
    "Capture /members with 6+ profile cards. Pick sort = Cards.",
    route="/members",
)

h2("Visiting another collector")
para(
    "Click any member card to open their profile. You'll see their bio, favorite team / "
    "players / chase list, favorite cards showcase, and any sets they've shared publicly. "
    "Their want list isn't visible — that's private — but if any of their shared sets overlap "
    "with what you're collecting, those overlaps show in the listings feed."
)

h2("Want-list hits feed")
para(
    "Already covered in Section 5. Worth re-emphasizing here: when another approved member "
    "lists or auctions a card on your want list, it surfaces in your home feed. This is the "
    "single biggest reason to keep your sets honest about what you don't own."
)

page_break()

# ───────────────────────── 11. ADMIN ─────────────────────────

h1("11. Admin")

para(
    "The Admin section is only visible to users with the admin flag on their profile. If "
    "you're a pilot user, you'll skip this section."
)

h2("Reviewing membership applications")
para(
    "/admin lists every account in pending status. Each row has the application form "
    "answers and an Approve / Decline button. Approving sends the welcome email and unlocks "
    "the rest of the site for that user."
)

h2("Set templates")
para(
    "The pre-loaded checklists every member sees in New upload come from /admin/templates. "
    "Upload a CSV with one row per card; the template becomes available across every "
    "member's account immediately."
)

h2("Set headers")
para(
    "/admin/set-headers controls the cover image and short description shown at the top of "
    "every set page (e.g. the iconic 1971 Topps black border, with a sentence about the set's "
    "history). Match by set slug; uploads accept JPEG/PNG/WebP."
)

page_break()

# ───────────────────────── 12. TROUBLESHOOTING ─────────────────────────

h1("12. Tips & troubleshooting")

h2("My set didn't import correctly")
bullets([
    "PSA CSVs change format occasionally — if a column doesn't map, open the file, confirm the headers match the standard PSA export, and re-upload.",
    "Make sure the file is saved as .csv (not .xlsx). Excel will let you save in either; pick CSV (Comma delimited).",
])

h2("eBay hits look stale")
para(
    "Click 🔄 Refresh on the eBay Hits panel — we cache aggressively (eBay's API has rate "
    "limits). A refresh forces a re-scan."
)

h2("My auction post copied with weird formatting")
para(
    "Facebook strips most formatting on paste. The system generates plain text on purpose. "
    "If you're seeing odd characters, check that your template doesn't have curly quotes "
    "(Word and Google Docs auto-convert these — paste your template body through a plain "
    "text editor first)."
)

h2("Where to send feedback or report a bug")
para(
    "During the pilot: email the admin directly with what you were doing, what you expected, "
    "and what happened. Screenshots help a lot. A more formal feedback channel is on the "
    "pilot launch checklist (Appendix B)."
)

h2("Data privacy & deletion")
para(
    "You own your data. Email the admin to request a full export or full deletion of your "
    "account; we honor deletion requests within 7 days. We don't sell or share member data."
)

page_break()

# ───────────────────────── APPENDIX A ─────────────────────────

h1("Appendix A — Keyboard shortcuts")

table = doc.add_table(rows=1, cols=2)
hdr = table.rows[0].cells
shade_cell(hdr[0], "3D1F4A"); shade_cell(hdr[1], "3D1F4A")
for i, txt in enumerate(["Shortcut", "Where it works"]):
    p = hdr[i].paragraphs[0]
    p.text = ''
    r = p.add_run(txt); r.bold = True; r.font.color.rgb = RGBColor(0xF8, 0xEC, 0xD0); r.font.size = Pt(11)

shortcuts = [
    ("← / →", "Set inventory lightbox — flip between cards"),
    ("Esc", "Close any modal (lightbox, picker, default-target)"),
    ("Cmd / Ctrl + K", "Global search (planned)"),
    ("Tab inside set table", "Move to the next editable cell"),
    ("Enter on a row checkbox", "Toggle bulk-select for that row"),
]
for short, where in shortcuts:
    row = table.add_row().cells
    p = row[0].paragraphs[0]; p.text = ''
    r = p.add_run(short); r.bold = True; r.font.size = Pt(10.5)
    row[1].text = where
    for c in row:
        for p_ in c.paragraphs:
            for run in p_.runs:
                run.font.size = Pt(10.5)
        set_cell_border(c, color="C8B8D0", size="4")

page_break()

# ───────────────────────── APPENDIX B ─────────────────────────

h1("Appendix B — Pilot launch checklist")

para(
    "Internal items to close before the broader pilot rollout. Not user-facing.",
    italic=True, color=INK_SOFT,
)

checklist = [
    ("Terms of Use / Disclaimer modal", "Liability + clear rules of the road. Must be acknowledged at first login."),
    ("Privacy policy page", "What we collect, how we use it. Required for any external users."),
    ("In-app feedback channel", "A Send feedback button that emails admin or creates a ticket so testers don't lose issues in DMs."),
    ("Pilot welcome email", "Extend the existing approval email with a Read this first link to this guide."),
    ("Quick-start checklist on home", "Dismissible card: Set your profile · Add your first set · Try a listing."),
    ("Analytics + error tracking", "Vercel Analytics + Sentry free tier covers a small pilot."),
    ("Backups confirmed", "Supabase paid plans include automatic backups — confirm before live data lands."),
    ("RLS policies double-checked", "Every table reviewed before non-admins arrive."),
    ("Structured feedback form", "Google Form or Typeform with screenshot upload."),
    ("Versioned changelog / What's new", "Tells testers when their feedback shipped."),
]
for i, (title, why) in enumerate(checklist, 1):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{title} — "); r.bold = True
    p.add_run(why)

page_break()

# ───────────────────────── COLOPHON ─────────────────────────

h1("Colophon")
para(
    "Sports Collective is a pilot project. This guide is generated from the codebase and "
    "updated as features ship — re-run docs/build_user_guide.py to regenerate.",
    italic=True, color=INK_SOFT,
)
para(
    f"Last regenerated: 2026-05-07.",
    italic=True, color=INK_SOFT, size=10,
)

# ───────────────────────── SAVE ─────────────────────────

out_path = Path(__file__).parent / "Sports-Collective-User-Guide.docx"
doc.save(out_path)
print(f"Wrote {out_path}")
